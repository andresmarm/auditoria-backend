import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_docx(plan) -> bytes:
    """
    Genera un archivo DOCX a partir del plan de auditoría guardado.
    Retorna los bytes del archivo listo para subir a Storage.
    """
    doc = Document()

    # Estilos generales
    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Calibri"
    estilo_normal.font.size = Pt(11)

    # === ENCABEZADO ===
    titulo = doc.add_heading("PLAN DE AUDITORÍA", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if plan.sesion:
        subtitulo = doc.add_paragraph(f"Proceso: {plan.sesion.nombre_proceso or 'N/A'}")
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Fecha de generación: {plan.created_at.strftime('%d/%m/%Y') if plan.created_at else 'N/A'}")
    doc.add_paragraph(f"Generado por: Sistema de Auditoría con IA")
    doc.add_paragraph("")

    # === CONTENIDO PRINCIPAL ===
    if plan.contenido_texto:
        _agregar_contenido_markdown(doc, plan.contenido_texto)

    # === NORMAS CITADAS ===
    if plan.normas_citadas:
        doc.add_heading("Normatividad de Referencia", level=1)
        tabla = doc.add_table(rows=1, cols=3)
        tabla.style = "Table Grid"

        encabezados = tabla.rows[0].cells
        encabezados[0].text = "Norma"
        encabezados[1].text = "Descripción"
        encabezados[2].text = "Artículo"

        for norma in plan.normas_citadas:
            fila = tabla.add_row().cells
            fila[0].text = norma.get("codigo", "")
            fila[1].text = norma.get("nombre", "")
            fila[2].text = norma.get("articulo", "")

    # === PIE ===
    doc.add_paragraph("")
    pie = doc.add_paragraph(
        "Documento generado automáticamente por el Asistente de Auditoría. "
        "Verifique la vigencia de la normatividad citada antes de aplicar."
    )
    pie.runs[0].font.size = Pt(9)
    pie.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Convertir a bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _agregar_contenido_markdown(doc: Document, texto: str):
    """
    Parsea el texto con formato markdown básico y lo agrega al documento.
    Detecta: ## Títulos, **negrita**, listas con -, tablas con |
    """
    lineas = texto.split("\n")
    i = 0

    while i < len(lineas):
        linea = lineas[i]

        # Títulos H1
        if linea.startswith("# "):
            doc.add_heading(linea[2:].strip(), level=1)

        # Títulos H2
        elif linea.startswith("## "):
            doc.add_heading(linea[3:].strip(), level=2)

        # Títulos H3
        elif linea.startswith("### "):
            doc.add_heading(linea[4:].strip(), level=3)

        # Listas
        elif linea.strip().startswith("- ") or linea.strip().startswith("* "):
            p = doc.add_paragraph(linea.strip()[2:], style="List Bullet")

        # Línea vacía
        elif not linea.strip():
            doc.add_paragraph("")

        # Párrafo normal
        else:
            p = doc.add_paragraph(linea)

        i += 1
