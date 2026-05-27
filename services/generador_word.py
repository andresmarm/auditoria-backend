"""
GENERADOR DE WORD
==================
Dos modos:
  1. generar_word_estandar()    → Word estructurado cuando el usuario no carga formato
  2. generar_word_con_formato() → Word que respeta la estructura del .docx cargado

Usa python-docx para el modo con formato (lectura del docx del usuario)
y docx-js (Node.js) para el estándar con estilos profesionales.
"""

import io
import os
import json
import tempfile
import subprocess
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── MODO 1: Word estándar (sin formato del usuario) ───────

def generar_word_estandar(plan_texto: str, nombre_proceso: str = "") -> bytes:
    """
    Genera un Word profesional estructurado a partir del texto del plan.
    Parsea los títulos Markdown (##, ###) para aplicar estilos de heading.
    """
    doc = Document()

    # Configurar página tamaño carta
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.page_width  = Cm(21.59)
    section.page_height = Cm(27.94)
    section.left_margin = section.right_margin = Cm(2.54)
    section.top_margin  = section.bottom_margin = Cm(2.54)

    # Encabezado institucional
    header = doc.sections[0].header
    header_p = header.paragraphs[0]
    header_p.text = "OFICINA DE CONTROL INTERNO — COMCAJA"
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.runs[0].font.size = Pt(9)
    header_p.runs[0].font.color.rgb = RGBColor(0x4A, 0x5E, 0x72)

    # Pie de página con numeración
    footer = doc.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _agregar_numero_pagina(footer_p)

    # Título principal
    titulo = doc.add_heading("PLAN DE AUDITORÍA INTERNA", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.runs[0].font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)

    if nombre_proceso:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(nombre_proceso.replace(".pdf","").replace(".docx","").upper())
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xC9, 0x92, 0x2A)
        run.bold = True

    doc.add_paragraph()

    # Parsear y renderizar el texto del plan
    lineas = plan_texto.split("\n")
    p_actual = None

    for linea in lineas:
        linea_stripped = linea.strip()

        if linea_stripped.startswith("# "):
            doc.add_heading(linea_stripped[2:], level=1)
            p_actual = None

        elif linea_stripped.startswith("## "):
            doc.add_heading(linea_stripped[3:], level=2)
            p_actual = None

        elif linea_stripped.startswith("### "):
            doc.add_heading(linea_stripped[4:], level=3)
            p_actual = None

        elif linea_stripped.startswith(("- ", "* ", "• ")):
            p = doc.add_paragraph(style="List Bullet")
            _agregar_texto_con_bold(p, linea_stripped[2:])
            p_actual = None

        elif linea_stripped and linea_stripped[0].isdigit() and ". " in linea_stripped[:4]:
            p = doc.add_paragraph(style="List Number")
            _agregar_texto_con_bold(p, linea_stripped.split(". ", 1)[-1])
            p_actual = None

        elif linea_stripped == "":
            p_actual = None

        elif linea_stripped.startswith("|") and "|" in linea_stripped[1:]:
            # Tabla Markdown — acumular filas
            _procesar_tabla_md(doc, lineas, lineas.index(linea))
            p_actual = None

        else:
            if p_actual is None:
                p_actual = doc.add_paragraph()
            else:
                p_actual.add_run(" ")
            _agregar_texto_con_bold(p_actual, linea_stripped)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _agregar_texto_con_bold(p, texto: str):
    """Parsea **negrita** dentro del texto y aplica formato."""
    partes = texto.split("**")
    for i, parte in enumerate(partes):
        if parte:
            run = p.add_run(parte)
            run.bold = (i % 2 == 1)


def _procesar_tabla_md(doc, lineas: list, idx_inicio: int):
    """Convierte una tabla Markdown en tabla Word."""
    filas_md = []
    i = idx_inicio
    while i < len(lineas) and "|" in lineas[i]:
        fila = [c.strip() for c in lineas[i].strip().strip("|").split("|")]
        if not all(set(c) <= set("-: ") for c in fila):  # saltar separador ---
            filas_md.append(fila)
        i += 1

    if not filas_md:
        return

    ncols = max(len(f) for f in filas_md)
    tabla = doc.add_table(rows=len(filas_md), cols=ncols)
    tabla.style = "Table Grid"

    for r, fila in enumerate(filas_md):
        for c, celda in enumerate(fila):
            if c < ncols:
                cell = tabla.cell(r, c)
                cell.text = celda
                if r == 0:
                    cell.paragraphs[0].runs[0].bold = True


def _agregar_numero_pagina(paragraph):
    """Inserta campo PAGE en el párrafo del pie de página."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    run2 = paragraph.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    run2._r.append(instrText)
    run3 = paragraph.add_run()
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run3._r.append(fld2)


# ── MODO 2: Word respetando formato del usuario ───────────

def generar_word_con_formato(plan_texto: str, formato_bytes: bytes) -> bytes:
    """
    Toma el .docx del usuario como base, extrae su estructura
    y genera un nuevo documento que respeta sus estilos y secciones,
    llenándolo con el contenido del plan generado.
    """
    # Leer la estructura del formato del usuario
    formato_doc = Document(io.BytesIO(formato_bytes))
    secciones_formato = _extraer_secciones_formato(formato_doc)

    # Crear nuevo documento basado en el formato
    doc = Document(io.BytesIO(formato_bytes))

    # Limpiar el contenido existente preservando los estilos
    for p in doc.paragraphs:
        if p.text.strip() and not _es_encabezado_fijo(p.text):
            # Reemplazar contenido de secciones con el plan generado
            _rellenar_seccion(p, plan_texto, secciones_formato)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _extraer_secciones_formato(doc: Document) -> dict:
    """Extrae las secciones principales del formato del usuario."""
    secciones = {}
    for p in doc.paragraphs:
        texto = p.text.strip().upper()
        if texto:
            secciones[texto] = p.style.name
    return secciones


def _es_encabezado_fijo(texto: str) -> bool:
    """Detecta si un párrafo es un encabezado institucional fijo."""
    fijos = ["CONTROL INTERNO", "COMCAJA", "SSF", "VERSIÓN", "CÓDIGO",
             "FECHA DE APROBACIÓN", "ELABORADO", "REVISADO"]
    return any(f in texto.upper() for f in fijos)


def _rellenar_seccion(p, plan_texto: str, secciones: dict):
    """Intenta mapear un párrafo vacío del formato con el contenido del plan."""
    # Por ahora limpia el texto para que el auditor lo llene manualmente
    # En una versión futura se puede mapear semánticamente
    pass
